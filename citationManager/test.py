from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK

document = Document("CID Fall Semester Final Report.docx")

# Iterate through sections and create new pages
for i, section in enumerate(document.sections):
    if i > 0:
        document.add_section(WD_SECTION.NEW_PAGE)

    # Count words in the section
    word_count = sum(
        len(paragraph.text.split())
        for paragraph in document.paragraphs
        if paragraph._element.getparent() == section._sectPr.getparent()
    )

    # Add word count as a section break
    section_break = document.add_paragraph()
    section_break.text = f"Page {i + 1} - Word Count: {word_count}"
    try:
        section_break.style = document.styles["Intense Quote"]
    except KeyError:
        # If 'Intense Quote' style doesn't exist, create it
        from docx.enum.style import WD_STYLE_TYPE

        document.styles.add_style("Intense Quote", WD_STYLE_TYPE.PARAGRAPH)
        section_break.style = document.styles["Intense Quote"]

    # Add a page break after the section break
    section_break.runs[0].add_break(WD_BREAK.PAGE)

# print(document.paragraphs[0].text)

document.save("document_with_different_footers.docx")

# ! does not make section at the end of every page...
# docx functionality does not work.
